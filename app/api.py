import json
import os
import pathlib
import datetime
import zipfile
import shutil
import re
import hashlib
import random
import string

import pandas
from docx import Document
import jwt
from PIL import Image, ImageFont, ImageDraw, ImageFilter

from common.log import get_logger
from config.setting import FILE_PATH, DatabaseTable, ITEM_LIST
from common.postgres_driver import create_conn, execute_sqls, execute_many, fetch_one, fetch_all
from app.item import DefaultItem
from app.cell import (
    CellRtempDchgCapacity,
    CellStandardCycleLife
)
from app.mod import (
    ModRtempDchgCapacity,
    ModSpecificPower,
    ModRtempRateDchgCapacity,
    ModReserve,
    ModRtempRateChgPerformance,
    ModHtempChargeMaintain,
    ModRtempChargeMaintain,
    ModLtempDchgCapacity,
    ModHtempDchgCapacity
)
from app.pack import (
    PackRtempDchgCapacity,
    PackEnergyDensity
)


def get_verify_code():
    def rndColor():
        return (random.randint(32, 127), random.randint(32, 127), random.randint(32, 127))

    def gene_text():
        return ''.join(random.sample(string.ascii_letters+string.digits, 4))

    def draw_lines(draw, num, width, height):
        for num in range(num):
            x1 = random.randint(0, width / 2)
            y1 = random.randint(0, height / 2)
            x2 = random.randint(0, width)
            y2 = random.randint(height / 2, height)
            draw.line(((x1, y1), (x2, y2)), fill='black', width=1)

    code = gene_text()

    width, height = 120, 50

    im = Image.new('RGB',(width, height),'white')

    font = ImageFont.truetype('static/arial.ttf', 40)

    draw = ImageDraw.Draw(im)

    for item in range(4):
        draw.text((5+random.randint(-3,3)+23*item, 5+random.randint(-3,3)),
                  text=code[item], fill=rndColor(),font=font )

    draw_lines(draw, 2, width, height)

    im = im.filter(ImageFilter.GaussianBlur(radius=1.0))

    return im, code


def add_user(user, password):
    """
    It takes a user name and password, hashes the password, and inserts the user name and hashed
    password into a database table
    
    :param user: the user name
    :param password: '123456'
    :return: The return value is the number of rows affected by the last SQL statement.
    """

    logger = get_logger()

    logger.info(f"call add_user: {user}")

    md5 = hashlib.md5()
    md5.update(password.encode('utf-8'))

    sql = f"INSERT INTO {DatabaseTable.USER}(usr,password) VALUES('{user}','{md5.hexdigest()}')"

    with create_conn() as conn:
        return execute_sqls(conn, sql)


def delete_user(user, password):
    """
    It deletes a user from the database if the user and password match
    
    :param user: the user name
    :param password: the password of the user to be deleted
    :return: The number of rows affected by the SQL statement.
    """

    logger = get_logger()

    logger.info(f"call delete_user: {user}")

    md5 = hashlib.md5()
    md5.update(password.encode('utf-8'))

    sql = f"DELETE FROM {DatabaseTable.USER} WHERE usr='{user}' AND password='{md5.hexdigest()}'"

    with create_conn() as conn:
        return execute_sqls(conn, sql)


def change_password(user, old_password, new_password):
    """
    Change password of user.
    
    :param user: The user name
    :param old_password: the old password
    :param new_password: the new password
    """

    logger = get_logger()

    logger.info(f"call change_password: {user}")

    sql = f"SELECT password FROM {DatabaseTable.USER} WHERE usr='{user}'"

    with create_conn() as conn:
        ret, data = fetch_one(conn, sql)

    if ret is False:
        return ret, data

    if data is None:
        return False, f"Can not find user {user}."

    old_md5 = hashlib.md5()
    old_md5.update(old_password.encode('utf-8'))

    if old_md5.hexdigest() != data[0]:
        return False, "old password is wrong."

    md5 = hashlib.md5()
    md5.update(new_password.encode('utf-8'))

    sql = f"UPDATE {DatabaseTable.USER} SET password='{md5.hexdigest()}' WHERE usr='{user}'"

    with create_conn() as conn:
        return execute_sqls(conn, sql)


def verify_user(user, password):
    """
    It takes a username and password, and returns a token if the password is correct
    
    :param user: The user name
    :param password: the password you want to verify
    """

    logger = get_logger()

    logger.info(f"call verify_user: {user}")

    sql = f"SELECT password FROM {DatabaseTable.USER} WHERE usr='{user}'"

    with create_conn() as conn:
        ret, data = fetch_one(conn, sql)

    if ret is False:
        return ret, data

    if data is None:
        return False, f"Can not find user {user}."

    md5 = hashlib.md5()
    md5.update(password.encode('utf-8'))

    if md5.hexdigest() != data[0]:
        return False, "password is wrong."

    return True, jwt.encode({'exp': (datetime.datetime.now() + datetime.timedelta(days=1)).timestamp(), "user": user}, 'catarc', algorithm='HS256')


def verify_token(token):
    """
    It takes a token, decodes it, and returns the user
    
    :param token: The token to be verified
    :return: True, params["user"]
    """

    logger = get_logger()

    logger.info(f"call verify_token: {token}")

    try:
        params = jwt.decode(token, 'catarc', algorithms=['HS256'])
    except Exception as e:
        return False, f"{e}"

    return True, params["user"]


def save_task_info(task_id, submit_date, submit_company, task_note, test_samples):
    """
    It takes in a task_id, submit_date, submit_company, task_note, and test_samples, and then it creates
    a directory with the task_id, and then it creates a sql statement to insert the data into the
    database
    
    :param task_id: The task ID, which is the name of the folder where the task files are stored
    :param submit_date: the date the task was submitted
    :param submit_company: The company that submitted the task
    :param task_note: The task note is a string that is used to describe the task
    :param test_samples: [{'sample_id': 'sample_id_1', 'sample_name': 'sample_name_1', 'sample_note':
    'sample_note_1'}, {'sample_id': 'sample_id_2', 'sample_name': 'sample_
    """

    logger = get_logger()

    logger.info(f"call save_task_info: task_id={task_id} submit_date={submit_date} submit_company={submit_company} task_note={task_note} test_samples={test_samples}")

    if submit_date == '':
        submit_date = datetime.datetime.now().strftime('%Y-%m-%d')

    data = {
        "task_id": task_id,
        "submit_date": submit_date,
        "submit_company": submit_company,
        "task_note": task_note,
        "test_samples": json.dumps(test_samples, ensure_ascii=False, separators=(',', ':'))
    }

    pathlib.Path(os.path.join(FILE_PATH, task_id)).mkdir(exist_ok=True)

    sql = f"SELECT task_id FROM {DatabaseTable.TASK} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, value = fetch_one(conn, sql)

    if ret is False:
        return ret, value

    if value is not None and value[0] is not None:
        values = ','.join([f"{k}='{v}'" for k, v in data.items() if k != "task_id"])

        sql = f"UPDATE {DatabaseTable.TASK} SET {values} WHERE task_id='{task_id}'"
    else:
        keys = ','.join(data.keys())
        values = ','.join([f"'{v}'" for v in data.values()])

        sql = f"INSERT INTO {DatabaseTable.TASK}({keys}) VALUES({values})"

    with create_conn() as conn:
        return execute_sqls(conn, sql)


def extract_product_file(task_id, product_file):
    """
    It takes a task_id and a product_file, and then it does a bunch of stuff with the product_file, and
    then it updates the database with the results of the stuff it did with the product_file
    
    :param task_id: the task id
    :param product_file: the name of the file to be processed
    """

    logger = get_logger()

    logger.info(f"call extract_product_file: task_id={task_id} product_file={product_file}")

    data = []
    doc_file = Document(os.path.join(FILE_PATH, task_id, "product_file" + os.path.splitext(product_file)[-1]))

    for tb in doc_file.tables:
        for i, row in enumerate(tb.rows):
            if row.cells[1].paragraphs[0].text == "":
                data.append(
                    [
                        "".join([para.text.strip() for para in row.cells[0].paragraphs]),
                        "".join([para.text.strip() for para in row.cells[2].paragraphs]),
                        "".join([para.text.strip() for para in row.cells[3].paragraphs])
                    ]
                )

    df = pandas.DataFrame(data, columns=["category", "name", "value"])
    df.insert(0, "task_id", task_id)

    sql = f"SELECT product_file FROM {DatabaseTable.TASK} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, value = fetch_one(conn, sql)

    if ret is False:
        return ret, value

    if value is not None and value[0] is not None:
        sql = f"DELETE FROM {DatabaseTable.PRODUCT} WHERE task_id='{task_id}'"

        with create_conn() as conn:
            ret, msg = execute_sqls(conn, sql)

        if ret is False:
            return ret, msg

        pathlib.Path(os.path.join(FILE_PATH, task_id, value[0])).unlink(missing_ok=True)

    sql = f"INSERT INTO {DatabaseTable.PRODUCT}({','.join(df.columns.tolist())}) VALUES({','.join(['%s'] * df.columns.size)})"

    with create_conn() as conn:
        ret, msg = execute_many(conn, sql, df.to_numpy().tolist())

    if ret is False:
        return ret, msg

    sql = f"UPDATE {DatabaseTable.TASK} SET product_date='{str(datetime.date.today())}', product_file='{product_file}' WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sql)

    if ret is False:
        return ret, msg

    pathlib.Path(os.path.join(FILE_PATH, task_id, "product_file" + os.path.splitext(product_file)[-1])).rename(pathlib.Path(os.path.join(FILE_PATH, task_id, product_file)))

    return True, data


def upload_test_file(user, task_id, test_file):
    """
    It takes a zip file, unzips it, renames the files, and then checks if the files are in the correct
    directory structure
    
    :param user: the user name
    :param task_id: the task id
    :param test_file: the name of the zip file
    """

    logger = get_logger()

    logger.info(f"call upload_test_file: user={user} task_id={task_id} test_file={test_file}")

    file_names = []
    test_dirs = []
    test_file_name = ""

    zip_file = zipfile.ZipFile(os.path.join(FILE_PATH, task_id, test_file), 'r')

    for old_name in zip_file.namelist():
        try:
            new_name = old_name.encode('cp437').decode('gbk')
        except:
            new_name = old_name.encode('utf-8').decode('utf-8')

        file_names.append(
            {
                "new": os.path.join(FILE_PATH, task_id, new_name),
                "old": os.path.join(FILE_PATH, task_id, old_name),
                "count": len(new_name.rstrip('/').split('/'))
            }
        )

        new_name_list = new_name.split('/')

        if test_file_name == "":
            test_file_name = new_name_list[0]

    if os.path.isfile(os.path.join(FILE_PATH, task_id, test_file_name)):
        pathlib.Path(os.path.join(FILE_PATH, task_id, test_file_name)).unlink()

    if os.path.isdir(os.path.join(FILE_PATH, task_id, test_file_name)):
        shutil.rmtree(os.path.join(FILE_PATH, task_id, test_file_name))

    zip_file.extractall(os.path.join(FILE_PATH, task_id))

    zip_file.close()

    pathlib.Path(os.path.join(FILE_PATH, task_id, test_file)).unlink()

    file_names.sort(key=lambda x: x["count"])

    try:
        for n in file_names:
            old_file = pathlib.Path(n["new"]).parent.joinpath(pathlib.Path(n["old"]).name)
            old_file.rename(pathlib.Path(n["new"]))

        if os.path.isfile(os.path.join(FILE_PATH, task_id, user)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, user)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, user)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, user))

        pathlib.Path(os.path.join(FILE_PATH, task_id, test_file_name)).rename(pathlib.Path(os.path.join(FILE_PATH, task_id, user)))
    except Exception as e:
        if os.path.isfile(os.path.join(FILE_PATH, task_id, user)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, user)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, user)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, user))

        if os.path.isfile(os.path.join(FILE_PATH, task_id, test_file_name)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, test_file_name)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, test_file_name)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, test_file_name))

        return False, f"Caught exception: {e.__doc__}({e})"

    with os.scandir(os.path.join(FILE_PATH, task_id, user)) as gbt_it:
        for gbt_entry in gbt_it:
            if gbt_entry.is_dir() and gbt_entry.name in ITEM_LIST:
                with os.scandir(gbt_entry.path) as category_it:
                    for category_entry in category_it:
                        if category_entry.is_dir() and category_entry.name in ITEM_LIST[gbt_entry.name]:
                            with os.scandir(category_entry.path) as item_it:
                                for item_entry in item_it:
                                    if item_entry.is_dir() and item_entry.name in ITEM_LIST[gbt_entry.name][category_entry.name]:
                                        test_dirs.append([gbt_entry.name, category_entry.name, item_entry.name])
                                    else:
                                        return False, f"{item_entry.path} is invalid."             
                        else:
                            return False, f"{category_entry.path} is invalid."            
            else:
                return False, f"{gbt_entry.path} is invalid."

    return True, test_dirs


def extract_test_file(user, task_id, test_items, parser_rule):
    """
    It takes a list of test items, and copies them from the user's directory to the task directory
    
    :param user: the user who is running the test
    :param task_id: the task id
    :param test_items: [('gbt', 'category', 'item', 'item_note')]
    :param parser_rule: the name of the parser rule
    """

    logger = get_logger()

    logger.info(f"call extract_test_file: user={user} task_id={task_id} test_items={test_items} parser_rule={parser_rule}")

    sqls = []

    for row in test_items:
        pathlib.Path(os.path.join(FILE_PATH, task_id, row[0], row[1], row[2])).mkdir(parents=True, exist_ok=True)

        shutil.copytree(os.path.join(FILE_PATH, task_id, user, row[0], row[1], row[2]), os.path.join(FILE_PATH, task_id, row[0], row[1], row[2]), dirs_exist_ok=True)

        sql = f"SELECT task_id FROM {DatabaseTable.STAT} WHERE task_id='{task_id}' AND gbt='{row[0]}' AND category='{row[1]}' AND item='{row[2]}'"

        with create_conn() as conn:
            ret, value = fetch_one(conn, sql)

        if ret is False:
            return ret, value

        if value is not None and value[0] is not None:
            sqls.append(f"UPDATE {DatabaseTable.STAT} SET status='extract',upload_user='{user}',item_note='{row[3]}',result=NULL WHERE task_id='{task_id}' AND gbt='{row[0]}' AND category='{row[1]}' AND item='{row[2]}'")
        else:
            sqls.append(f"INSERT INTO {DatabaseTable.STAT}(task_id,status,upload_user,item_note,gbt,category,item) VALUES('{task_id}','extract','{user}','{row[3]}','{row[0]}','{row[1]}','{row[2]}')")

    if not sqls:
        return True, ""

    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sqls)

    if ret is False:
        return ret, msg

    sql = f"UPDATE {DatabaseTable.TASK} SET test_date='{str(datetime.date.today())}', test_parser='{parser_rule}' WHERE task_id='{task_id}' AND test_date IS NULL"

    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sql)

    if ret is False:
        return ret, msg

    return True, ""


def delete_task(task_id):
    """
    It deletes a task from the database and deletes the task's folder from the file system
    
    :param task_id: the task id
    """

    logger = get_logger()

    logger.info(f"call delete_task: task_id={task_id}")

    sqls = [
        f"DELETE FROM {DatabaseTable.STAT} WHERE task_id='{task_id}'",
        f"DELETE FROM {DatabaseTable.PRODUCT} WHERE task_id='{task_id}'",
        f"DELETE FROM {DatabaseTable.TASK} WHERE task_id='{task_id}'"
    ]


    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sqls)

    if ret is False:
        return ret, msg

    try:
        shutil.rmtree(os.path.join(FILE_PATH, task_id))
    except:
        pass

    return True, ""


def query_task_upload_info(task_id, task_note, gbt, item, item_note, test_sample_company, test_sample_name, product_date_range, test_date_range):
    """
    This function queries the database for the task_id, task_note, gbt, item, item_note,
    test_sample_company, test_sample_name, product_date_range, and test_date_range for the given task_id
    
    :param task_id: the task id
    :param task_note: the note of the task
    :param gbt: the GBT number of the test
    :param item: the item name
    :param item_note: the note of the item
    :param test_sample_company: the company that provided the test sample
    :param test_sample_name: the name of the test sample
    :param product_date_range: a list of two strings, the first is the start date, the second is the end
    date
    :param test_date_range: a list of two strings, the first one is the start date, the second one is
    the end date
    """

    logger = get_logger()

    logger.info(f"call query_task_upload_info: task_id={task_id} task_note={task_note} gbt={gbt} item={item} item_note={item_note} test_sample_company={test_sample_company} test_sample_name={test_sample_name} product_date_range={product_date_range} test_date_range={test_date_range}")

    sql = f"SELECT task_id, task_note, test_samples, product_date, test_date FROM {DatabaseTable.TASK}"

    if task_id != "" or task_note != "" or product_date_range["from"] != "" or test_date_range["from"] != "":
        sql += " WHERE"

        if task_id != "":
            sql += f" task_id like '%{task_id}%'"

        if task_note != "":
            if not sql.endswith("WHERE"):
                sql += " AND"

            sql += f" task_note like '%{task_note}%'"

        if product_date_range["from"] != "":
            from_date = product_date_range["from"]
            to_date = product_date_range["to"]

            if not sql.endswith("WHERE"):
                sql += " AND"

            sql += f" product_date >= '{from_date}' AND product_date <= '{to_date}'"

        if test_date_range["from"] != "":
            from_date = test_date_range["from"]
            to_date = test_date_range["to"]

            if not sql.endswith("WHERE"):
                sql += " AND"

            sql += f" test_date >= '{from_date}' AND test_date <= '{to_date}'"

    sql += " ORDER BY create_time DESC"

    with create_conn() as conn:
        ret, data = fetch_all(conn, sql)

    if ret is False:
        return ret, data

    task_list = []

    for task in data:
        company = None
        name = None

        samples = json.loads(task[2])

        if test_sample_company != "":
            for sample in samples[:]:
                if sample["company"].find(test_sample_company) >= 0:
                    company = sample["company"]
                    break
        else:
            if len(samples) > 0:
                company = samples[0]["company"]
            else:
                company = ""

        if test_sample_name != "":
            for sample in samples[:]:
                if sample["name"].find(test_sample_name) >= 0:
                    name = sample["name"]
                    break
        else:
            if len(samples) > 0:
                name = samples[0]["name"]
            else:
                name = ""

        if company is not None and name is not None:
            if gbt != "" or item != "" or item_note != "":
                sql = f"SELECT task_id FROM {DatabaseTable.STAT} WHERE task_id = '{task[0]}'"

                if gbt != "":
                    sql += f" AND gbt LIKE '%{gbt}%'"

                if item != "":
                    sql += f" AND item LIKE '%{item}%'"                    

                if item_note != "":
                    sql += f" AND item_note LIKE '%{item_note}%'" 

                with create_conn() as conn:
                    ret, value = fetch_one(conn, sql)

                if ret is False:
                    return ret, value

                if value is None or value[0] is None:
                    continue

            task_list.append([task[0], company, task[3].isoformat() if task[3] is not None else "", task[4].isoformat() if task[4] is not None else ""])

    return True, task_list


def query_task_product_info(task_id):
    """
    It queries the database for the product information of a task
    
    :param task_id: the task id
    :return: A list of tuples.
    """

    logger = get_logger()

    logger.info(f"call query_task_product_info: task_id={task_id}")

    sql = f"SELECT category, name, value FROM {DatabaseTable.PRODUCT} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        return fetch_all(conn, sql)


def query_task_product_file(task_id):
    """
    It queries the database for a file name, and returns the file name if it exists
    
    :param task_id: The task ID
    """

    logger = get_logger()

    logger.info(f"call query_task_product_file: task_id={task_id}")

    sql = f"SELECT product_file FROM {DatabaseTable.TASK} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, value = fetch_one(conn, sql)

    if ret is False:
        return ret, value

    if value is None or value[0] is None:
        return False, f"Can not find task {task_id}."

    return True, value[0]


def query_task_test_info(task_id):
    """
    It queries the database for a task's test information
    
    :param task_id: task_id
    """

    logger = get_logger()

    logger.info(f"call query_task_test_info: task_id={task_id}")

    sql = f"SELECT submit_date, submit_company, task_note, test_samples, test_date FROM {DatabaseTable.TASK} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, data = fetch_one(conn, sql)

    if ret is False:
        return ret, data

    if data is None:
        return False, f"Can not find task {task_id}."

    sql = f"SELECT gbt, category, item, item_note, status, upload_user FROM {DatabaseTable.STAT} WHERE task_id='{task_id}'"

    with create_conn() as conn:
        ret, items = fetch_all(conn, sql)

    if ret is False:
        return ret, items

    if items is None:
        return False, f"Can not find task {task_id} items."

    items.sort(key=lambda x: x[2])
    items.sort(key=lambda x: x[1])
    items.sort(key=lambda x: x[0])

    return True, {"submit_date": data[0].isoformat(), "submit_company": data[1], "task_note": data[2], "test_samples": json.loads(data[3]), "test_date": data[4].isoformat() if data[4] is not None else "", "test_items": items}


def extract_item_file(user, task_id, gbt, category, item, item_file, item_note):
    """
    It takes a zip file, extracts it, and then moves the contents of the zip file to a new directory
    
    :param user: the user who uploaded the file
    :param task_id: the task id
    :param gbt: the name of the project
    :param category: SampleCategory.CELL
    :param item: the name of the item, such as "cell_1"
    :param item_file: the file name of the zip file
    :param item_note: the note of the item
    """

    logger = get_logger()

    logger.info(f"call extract_item_file: user={user} task_id={task_id} gbt={gbt} category={category} item={item} item_file={item_file} item_note={item_note}")

    if gbt in ITEM_LIST:
        if category in ITEM_LIST[gbt]:
            if item not in ITEM_LIST[gbt][category]:
                return False, f"{item} is invalid."             
        else:
            return False, f"{category} is invalid."            
    else:
        return False, f"{gbt} is invalid."

    file_names = []
    samples_dirs = []
    item_file_name = ""

    zip_file = zipfile.ZipFile(os.path.join(FILE_PATH, task_id, item_file), 'r')

    for old_name in zip_file.namelist():
        try:
            new_name = old_name.encode('cp437').decode('gbk')
        except:
            new_name = old_name.encode('utf-8').decode('utf-8')

        file_names.append(
            {
                "new": os.path.join(FILE_PATH, task_id, new_name),
                "old": os.path.join(FILE_PATH, task_id, old_name),
                "count": len(new_name.rstrip('/').split('/'))
            }
        )

        new_name_list = new_name.split('/')

        if item_file_name == "":
            item_file_name = new_name_list[0]

    if os.path.isfile(os.path.join(FILE_PATH, task_id, item_file_name)):
        pathlib.Path(os.path.join(FILE_PATH, task_id, item_file_name)).unlink()

    if os.path.isdir(os.path.join(FILE_PATH, task_id, item_file_name)):
        shutil.rmtree(os.path.join(FILE_PATH, task_id, item_file_name))

    zip_file.extractall(os.path.join(FILE_PATH, task_id))

    zip_file.close()

    pathlib.Path(os.path.join(FILE_PATH, task_id, item_file)).unlink()

    file_names.sort(key=lambda x: x["count"])

    try:
        for n in file_names:
            old_file = pathlib.Path(n["new"]).parent.joinpath(pathlib.Path(n["old"]).name)
            old_file.rename(pathlib.Path(n["new"]))

        if os.path.isfile(os.path.join(FILE_PATH, task_id, user)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, user)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, user)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, user))

        pathlib.Path(os.path.join(FILE_PATH, task_id, item_file_name)).rename(pathlib.Path(os.path.join(FILE_PATH, task_id, user)))
    except Exception as e:
        if os.path.isfile(os.path.join(FILE_PATH, task_id, user)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, user)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, user)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, user))

        if os.path.isfile(os.path.join(FILE_PATH, task_id, item_file_name)):
            pathlib.Path(os.path.join(FILE_PATH, task_id, item_file_name)).unlink()

        if os.path.isdir(os.path.join(FILE_PATH, task_id, item_file_name)):
            shutil.rmtree(os.path.join(FILE_PATH, task_id, item_file_name))

        return False, f"Caught exception: {e.__doc__}({e})"

    with os.scandir(os.path.join(FILE_PATH, task_id, user)) as sample_it:
        for sample_entry in sample_it:
            if sample_entry.name in ITEM_LIST:
                return False, f"{sample_entry.name} is invalid."

            samples_dirs.append(sample_entry.name)

    pathlib.Path(os.path.join(FILE_PATH, task_id, gbt, category, item)).mkdir(parents=True, exist_ok=True)

    shutil.copytree(os.path.join(FILE_PATH, task_id, user), os.path.join(FILE_PATH, task_id, gbt, category, item), dirs_exist_ok=True)

    sql = f"SELECT task_id FROM {DatabaseTable.STAT} WHERE task_id='{task_id}' AND gbt='{gbt}' AND category='{category}' AND item='{item}'"

    with create_conn() as conn:
        ret, value = fetch_one(conn, sql)

    if ret is False:
        return ret, value

    if value is not None and value[0] is not None:
        if item_note is None:
            sql = f"UPDATE {DatabaseTable.STAT} SET status='extract',upload_user='{user}',result=NULL WHERE task_id='{task_id}' AND gbt='{gbt}' AND category='{category}' AND item='{item}'"
        else:
            sql = f"UPDATE {DatabaseTable.STAT} SET status='extract',upload_user='{user}',item_note='{item_note}',result=NULL WHERE task_id='{task_id}' AND gbt='{gbt}' AND category='{category}' AND item='{item}'"
    else:
        if item_note is None:
            sql = f"INSERT INTO {DatabaseTable.STAT}(task_id,status,upload_user,gbt,category,item) VALUES('{task_id}','extract','{user}','{gbt}','{category}','{item}')"
        else:
            sql = f"INSERT INTO {DatabaseTable.STAT}(task_id,status,upload_user,item_note,gbt,category,item) VALUES('{task_id}','extract','{user}','{item_note}','{gbt}','{category}','{item}')"

    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sql)

    if ret is False:
        return ret, msg

    return True, samples_dirs


def parse_test_item(task_id, gbt, category, item):
    """
    It takes a task ID, a GBT, a category, and an item, and returns a dictionary with the item's ID, the
    task ID, the category, and the item's text
    
    :param task_id: the id of the task
    :param gbt: the GBT object
    :param category: the category of the item, e.g. "test_item_1"
    :param item: the item to be parsed
    """

    logger = get_logger()

    logger.info(f"call parse_test_item: task_id={task_id} gbt={gbt} category={category} item={item}")

    if gbt in ITEM_LIST:
        if category in ITEM_LIST[gbt]:
            if item in ITEM_LIST[gbt][category]:
                instance = eval(ITEM_LIST[gbt][category][item])(task_id, gbt, category, item)
            else:
                return False, f"{item} is invalid."             
        else:
            return False, f"{category} is invalid."            
    else:
        return False, f"{gbt} is invalid."

    ret, item_result = instance.preprocess()
    if ret is False:
        return ret, item_result

    logger.info(f"parse: {gbt} {category} {item}")

    if item_result:
        sql = f"UPDATE {DatabaseTable.STAT} SET status='parse', result='{json.dumps(item_result, separators=(',', ':'))}' WHERE task_id='{task_id}' AND gbt='{gbt}' AND category='{category}' AND item='{item}'"

        with create_conn() as conn:
            return execute_sqls(conn, sql)

        return True, "parse"

    return True, "extract"


def delete_test_item(task_id, gbt, category, item):
    """
    > Delete a test item from the database
    
    :param task_id: The task ID of the task you want to delete the item from
    :param gbt: the name of the GBT file
    :param category: the category of the item you want to delete
    :param item: the item to delete
    """

    logger = get_logger()

    logger.info(f"call delete_test_item: task_id={task_id} gbt={gbt} category={category} item={item}")

    sql = f"DELETE FROM {DatabaseTable.STAT} WHERE task_id='{task_id}' AND gbt='{gbt}' AND category='{category}' AND item='{item}'"

    with create_conn() as conn:
        ret, msg = execute_sqls(conn, sql)

    if ret is False:
        return ret, msg

    try:
        shutil.rmtree(os.path.join(FILE_PATH, task_id, gbt, category, item))
    except:
        pass

    return True, ""


def is_test_item_exist(task_id, gbt, category, item):
    """
    > This function checks if a test item exists in the database
    
    :param task_id: the task id of the task you want to check
    :param gbt: the name of the GBT
    :param category: the category of the item, e.g. "text", "image", "audio", "video"
    :param item: the item to be tested
    """

    logger = get_logger()

    logger.info(f"call is_test_item_exist: task_id={task_id} gbt={gbt} category={category} item={item}")

    if gbt in ITEM_LIST:
        if category in ITEM_LIST[gbt]:
            if item not in ITEM_LIST[gbt][category]:
                return False, f"{item} is invalid."             
        else:
            return False, f"{category} is invalid."            
    else:
        return False, f"{gbt} is invalid."

    if os.path.isdir(os.path.join(FILE_PATH, task_id, gbt, category, item)):
        return True, "yes"

    return True, "no"


def query_extract_item_info(task_id, gbt, category, item):
    """
    It takes in a task_id, gbt, category, and item, and returns a dictionary with two keys: files and
    decision
    
    :param task_id: the task id
    :param gbt: the name of the GBT
    :param category: SampleCategory.CELL
    :param item: the name of the item to be extracted
    """

    logger = get_logger()

    logger.info(f"call query_extract_item_info: task_id={task_id} gbt={gbt} category={category} item={item}")

    if gbt in ITEM_LIST:
        if category in ITEM_LIST[gbt]:
            if item in ITEM_LIST[gbt][category]:
                instance = eval(ITEM_LIST[gbt][category][item])(task_id, gbt, category, item)
            else:
                return False, f"{item} is invalid."             
        else:
            return False, f"{category} is invalid."            
    else:
        return False, f"{gbt} is invalid."

    ret, stat_data = instance.get_stat()
    if ret is False:
        return ret, stat_data

    ret, decision_data = instance.get_decision()
    if ret is False:
        return ret, decision_data

    return True, {"files": stat_data["files"], "decision": decision_data}


def query_parse_item_info(task_id, gbt, category, item):
    """
    It takes in a task_id, gbt, category, and item, and returns a dictionary of files, table, list,
    graph, and decision.
    
    :param task_id: the task id
    :param gbt: the name of the data source
    :param category: "category"
    :param item: the name of the item to be parsed
    """

    logger = get_logger()

    logger.info(f"call query_parse_item_info: task_id={task_id} gbt={gbt} category={category} item={item}")

    if gbt in ITEM_LIST:
        if category in ITEM_LIST[gbt]:
            if item in ITEM_LIST[gbt][category]:
                instance = eval(ITEM_LIST[gbt][category][item])(task_id, gbt, category, item)
            else:
                return False, f"{item} is invalid."             
        else:
            return False, f"{category} is invalid."            
    else:
        return False, f"{gbt} is invalid."

    ret, stat_data = instance.get_stat()
    if ret is False:
        return ret, stat_data

    ret, decision_data = instance.get_decision()
    if ret is False:
        return ret, decision_data

    return True, {"files": stat_data["files"], "table": stat_data["table"], "list": stat_data["list"], "graph": stat_data["graph"], "decision": decision_data}


def query_item_list(gbt, category):
    """
    > This function returns a list of items for a given category
    
    :param gbt: the name of the group
    :param category: The category of the item
    :return: A list of items
    """

    logger = get_logger()

    logger.info(f"call query_item_list: gbt={gbt} category={category}")

    if gbt in ITEM_LIST and category in ITEM_LIST[gbt]:
        return True, list(ITEM_LIST[gbt][category].keys())

    return True, []